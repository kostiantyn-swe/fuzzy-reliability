"""Convert TECHNICAL_DESCRIPTION.md to TECHNICAL_DESCRIPTION.docx.

Supports headings H1-H4, bold (**text**), inline code (`text`),
code blocks (```...```), tables (|col|col|), horizontal rules (---),
unordered lists (- item), and plain paragraphs.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(__file__).parent.parent / "TECHNICAL_DESCRIPTION.md"
DST = Path(__file__).parent.parent / "TECHNICAL_DESCRIPTION.docx"


# ── Стилі ────────────────────────────────────────────────────────────────────

def set_code_font(run, size_pt=9):
    run.font.name = "Courier New"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)


def shade_paragraph(para, fill="F0F0F0"):
    """Сірий фон для блоку коду."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_table_borders(tbl):
    """Тонкі межі таблиці."""
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ── Inline-форматування ───────────────────────────────────────────────────────

INLINE_PATTERN = re.compile(r"\*\*(.+?)\*\*|`(.+?)`")


def add_inline(para, text: str, default_bold=False, code_size=9):
    """Додає текст у параграф з підтримкою **bold** і `code`."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            if default_bold:
                run.bold = True
        if m.group(1) is not None:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            run = para.add_run(m.group(2))
            set_code_font(run, code_size)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        if default_bold:
            run.bold = True


# ── Таблиці ───────────────────────────────────────────────────────────────────

def collect_table(lines, start):
    """Збирає рядки таблиці починаючи з start. Повертає (rows, next_idx)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Розбити на комірки
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    # Видалити рядок-роздільник (---|---) якщо є
    if len(rows) >= 2 and all(re.match(r"^[-:]+$", c) for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    return rows, i


def render_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    add_table_borders(tbl)
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = tbl.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            para.style = doc.styles["Normal"]
            add_inline(para, cell_text, default_bold=(ri == 0))
            if ri == 0:
                # Фон заголовка таблиці
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DCE6F1")
                tcPr.append(shd)


# ── Головна функція конвертації ───────────────────────────────────────────────

def convert(src: Path, dst: Path):
    doc = Document()

    # Поля сторінки
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.0)

    # Базовий шрифт
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # Налаштування стилів заголовків
    heading_sizes = {1: 18, 2: 15, 3: 13, 4: 12}
    for lvl, size in heading_sizes.items():
        style = doc.styles[f"Heading {lvl}"]
        style.font.size = Pt(size)
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Горизонтальна риска → просто пропуск
        if re.match(r"^-{3,}$", stripped) or re.match(r"^_{3,}$", stripped):
            i += 1
            continue

        # Порожній рядок
        if not stripped:
            i += 1
            continue

        # Заголовки H1–H4
        m = re.match(r"^(#{1,4})\s+(.+)", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            para = doc.add_heading(level=level)
            para.clear()
            add_inline(para, text)
            i += 1
            continue

        # Блок коду
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            for cl in code_lines:
                p = doc.add_paragraph()
                shade_paragraph(p)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                run = p.add_run(cl)
                set_code_font(run)
            doc.add_paragraph()  # відступ після блоку
            continue

        # Таблиця
        if stripped.startswith("|"):
            rows, i = collect_table(lines, i)
            render_table(doc, rows)
            doc.add_paragraph()
            continue

        # Невпорядкований список
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines):
                s = lines[i].strip()
                m2 = re.match(r"^[-*]\s+(.+)", s)
                if m2:
                    items.append(m2.group(1))
                    i += 1
                elif not s:
                    i += 1
                    break
                else:
                    break
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                add_inline(p, item)
            continue

        # Звичайний параграф
        para = doc.add_paragraph()
        add_inline(para, stripped)
        i += 1

    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == "__main__":
    convert(SRC, DST)
